import PyPDF2
import sys
import docx
from docx import Document
from docx.shared import Pt


class PdfToDoc:
    def __init__(self, result):
        print("result "+ result)
        fileName = "E:/PycharmProjects/testMaterial/" + result
        print(fileName)
        pdfFileObj = open(fileName, 'rb')
        document = Document()

        pdfReader= PyPDF2.PdfFileReader(pdfFileObj)
        print(pdfReader.numPages)
        for i in range(0,pdfReader.numPages):
            pageObj = pdfReader.getPage(i)
            text = pageObj.extractText()
            print(text)

            p=document.add_paragraph().add_run(text)
            font = p.font
            font.name='OpenDyslexic'
            font.size=Pt(16)

            document.save('demo1.docx')

        pdfFileObj.close()